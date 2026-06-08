# ═══════════════════════════════════════════════════════════════════
#  ResumeIQ — Document Extractor (Python)
#  File: backend/doc_extractor.py
#
#  Handles text extraction from uploaded resume files:
#    .txt  → direct read
#    .pdf  → PyPDF2 full page extraction
#    .docx → python-docx paragraph extraction
#
#  This replaces the browser-side extraction (pdf.js, mammoth.js)
#  with server-side Python extraction — more robust and reliable.
#
#  [ML-HOOK] Future enhancements:
#    - OCR for scanned PDFs (Tesseract + pytesseract)
#    - Layout-aware extraction (pdfplumber, camelot for tables)
#    - Resume section parser (header/skills/experience detection)
# ═══════════════════════════════════════════════════════════════════

import os
from typing import Optional
from io import BytesIO


def extract_text(file_path: str = None, file_bytes: bytes = None,
                 filename: str = None) -> str:
    """
    Extract text from a resume file.

    Parameters:
        file_path (str): Path to the file on disk (for CLI use)
        file_bytes (bytes): Raw file bytes (for API upload)
        filename (str): Original filename (needed when using file_bytes
                        to determine file type)

    Returns:
        str: Extracted text content

    Raises:
        ValueError: If file type is unsupported
    """
    # Determine file extension
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
    elif filename:
        ext = os.path.splitext(filename)[1].lower()
    else:
        raise ValueError("Either file_path or filename must be provided")

    # Route to appropriate extractor
    if ext == '.txt':
        return _extract_txt(file_path, file_bytes)
    elif ext == '.pdf':
        return _extract_pdf(file_path, file_bytes)
    elif ext in ('.docx', '.doc'):
        return _extract_docx(file_path, file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: .txt, .pdf, .docx, .doc"
        )


# ═══════════════════════════════════════════════════════════════════
#  TXT Extraction — Simple file read
# ═══════════════════════════════════════════════════════════════════
def _extract_txt(file_path: str = None, file_bytes: bytes = None) -> str:
    """Extract text from a .txt file."""
    if file_bytes:
        return file_bytes.decode('utf-8', errors='ignore').strip()
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read().strip()


# ═══════════════════════════════════════════════════════════════════
#  PDF Extraction — Uses PyPDF2
#
#  Extracts text from every page of the PDF.
#  Works with all text-based PDFs. For scanned/image PDFs,
#  you'd need OCR (pytesseract).
#
#  [ML-HOOK] For production, consider:
#    import pytesseract
#    from pdf2image import convert_from_bytes
#    images = convert_from_bytes(pdf_bytes)
#    text = '\n'.join(pytesseract.image_to_string(img) for img in images)
# ═══════════════════════════════════════════════════════════════════
def _extract_pdf(file_path: str = None, file_bytes: bytes = None) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        if file_bytes:
            reader = PdfReader(BytesIO(file_bytes))
        else:
            reader = PdfReader(file_path)

        full_text = ''
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += page_text + '\n'

        result = full_text.strip()
        if len(result) < 50:
            print(f"  ⚠️  PDF extraction got very little text ({len(result)} chars) "
                  f"— might be a scanned/image PDF.")

        return result

    except ImportError:
        raise ImportError(
            "PyPDF2 is required for PDF extraction. "
            "Install it: pip install PyPDF2"
        )
    except Exception as e:
        print(f"  ❌ PDF extraction failed: {e}")
        return ''


# ═══════════════════════════════════════════════════════════════════
#  DOCX Extraction — Uses python-docx
#
#  .docx files are ZIP archives containing XML.
#  python-docx handles the unzipping and XML parsing automatically.
#
#  [ML-HOOK] For more advanced extraction:
#    - Extract tables separately for structured data
#    - Parse headers/footers for name/contact info
#    - Detect formatting (bold = skills, bullets = experience)
# ═══════════════════════════════════════════════════════════════════
def _extract_docx(file_path: str = None, file_bytes: bytes = None) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document

        if file_bytes:
            doc = Document(BytesIO(file_bytes))
        else:
            doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result = '\n'.join(paragraphs).strip()

        if len(result) < 50:
            print(f"  ⚠️  DOCX extraction got very little text ({len(result)} chars).")

        return result

    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install it: pip install python-docx"
        )
    except Exception as e:
        print(f"  ❌ DOCX extraction failed: {e}")
        return ''


# ===================================================================
#  SELF-TEST
# ===================================================================
if __name__ == '__main__':
    print("=" * 60)
    print("  ResumeIQ Document Extractor -- Self-Test")
    print("=" * 60)

    # Test TXT extraction from bytes
    test_bytes = b"Stuti Jain - CS + AI/ML Student\nSkills: Python, TensorFlow"
    text = extract_text(file_bytes=test_bytes, filename="test.txt")
    print(f"\n  TXT extraction: {len(text)} chars")
    print(f"  Content: {text[:80]}...")

    print(f"\n{'=' * 60}")
    print("  [OK] Document extractor ready!")
    print(f"{'=' * 60}\n")
